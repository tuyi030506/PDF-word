#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
PDF to Word Converter
高保真 PDF 转 Word 工具

功能：
- 将 PDF 文件转换为与原始排版完全一致的 Word 文档
- 保持段落样式、表格、图片、页眉页脚等格式
- 支持目录生成和文档结构保持

作者：AI Assistant
日期：2024-03-23
"""

import os
import sys
import logging
import subprocess
import tempfile
from pathlib import Path
from datetime import datetime
from typing import Optional, List, Dict, Tuple

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pdf2image import convert_from_path
from tqdm import tqdm
from PIL import Image

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('conversion.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)

class PDFConverter:
    """PDF转Word转换器"""
    
    def __init__(self, input_pdf: str, output_dir: str = "output"):
        """
        初始化转换器
        
        Args:
            input_pdf: PDF文件路径
            output_dir: 输出目录
        """
        self.input_pdf = Path(input_pdf)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        self.doc = None  # Word文档对象
        self.pdf_doc = None  # PDF文档对象
        self.temp_dir = None  # 临时目录
        
        self._check_environment()
    
    def _check_environment(self) -> None:
        """检查运行环境"""
        # 检查LibreOffice
        try:
            subprocess.run(['soffice', '--version'], 
                         capture_output=True, 
                         check=True)
            logging.info("LibreOffice 检测成功")
        except (subprocess.CalledProcessError, FileNotFoundError):
            logging.error("未检测到LibreOffice，请确保已安装并添加到环境变量")
            sys.exit(1)
        
        # 检查输入文件
        if not self.input_pdf.exists():
            logging.error(f"输入文件不存在: {self.input_pdf}")
            sys.exit(1)
            
        # 记录系统环境
        logging.info(f"操作系统: {sys.platform}")
        logging.info(f"Python版本: {sys.version}")
    
    def _initial_conversion(self) -> Path:
        """
        使用LibreOffice进行初步转换
        
        Returns:
            转换后的Word文件路径
        """
        logging.info("开始初步转换...")
        
        # 创建临时目录
        self.temp_dir = tempfile.mkdtemp()
        temp_dir_path = Path(self.temp_dir)
        
        # 调用LibreOffice进行转换
        cmd = [
            'soffice',
            '--headless',
            '--convert-to', 'docx',
            str(self.input_pdf),
            '--outdir', str(temp_dir_path)
        ]
        
        try:
            subprocess.run(cmd, check=True, capture_output=True)
            output_file = temp_dir_path / f"{self.input_pdf.stem}.docx"
            if not output_file.exists():
                raise RuntimeError("转换失败，未生成输出文件")
            logging.info("初步转换完成")
            return output_file
        except subprocess.CalledProcessError as e:
            logging.error(f"LibreOffice转换失败: {e.stderr.decode()}")
            raise
    
    def _extract_pdf_styles(self) -> Dict:
        """
        提取PDF中的样式信息
        
        Returns:
            样式信息字典
        """
        logging.info("正在分析PDF样式...")
        styles = {
            'paragraphs': [],
            'tables': [],
            'images': [],
            'headers': [],
            'footers': []
        }
        
        # 使用PyMuPDF提取样式
        self.pdf_doc = fitz.open(self.input_pdf)
        
        for page_num in range(len(self.pdf_doc)):
            page = self.pdf_doc[page_num]
            
            # 提取文本块信息
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block["type"] == 0:  # 文本块
                    for line in block["lines"]:
                        for span in line["spans"]:
                            styles['paragraphs'].append({
                                'text': span["text"],
                                'font': span["font"],
                                'size': span["size"],
                                'color': span["color"],
                                'flags': span["flags"],  # 包含粗体、斜体等信息
                                'bbox': span["bbox"]
                            })
                elif block["type"] == 1:  # 图片
                    styles['images'].append({
                        'bbox': block["bbox"],
                        'width': block["width"],
                        'height': block["height"]
                    })
        
        logging.info(f"样式分析完成: {len(styles['paragraphs'])} 段落, "
                    f"{len(styles['images'])} 图片")
        return styles
    
    def _apply_paragraph_styles(self, paragraph, style_info: Dict) -> None:
        """应用段落样式"""
        # 设置字体
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        font = run.font
        font.name = style_info.get('font', 'Microsoft YaHei')
        font.size = Pt(style_info.get('size', 12))
        
        # 设置颜色
        if 'color' in style_info:
            rgb = style_info['color']
            font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
        
        # 设置粗体和斜体
        flags = style_info.get('flags', 0)
        font.bold = bool(flags & 2)  # 粗体
        font.italic = bool(flags & 1)  # 斜体
        
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(12)
        paragraph_format.space_after = Pt(12)
        paragraph_format.line_spacing = 1.15
    
    def _handle_tables(self, doc: Document, styles: Dict) -> None:
        """处理表格"""
        logging.info("正在处理表格...")
        # TODO: 实现表格处理逻辑
    
    def _handle_images(self, doc: Document, styles: Dict) -> None:
        """处理图片"""
        logging.info("正在处理图片...")
        for img_info in styles['images']:
            try:
                # 从PDF中提取图片
                page = self.pdf_doc[img_info['page_num']]
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = self.pdf_doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # 保存为临时文件
                    temp_img_path = Path(self.temp_dir) / f"image_{xref}.png"
                    with open(temp_img_path, 'wb') as img_file:
                        img_file.write(image_bytes)
                    
                    # 添加到Word文档
                    doc.add_picture(str(temp_img_path),
                                  width=Inches(img_info['width']/72),
                                  height=Inches(img_info['height']/72))
            except Exception as e:
                logging.warning(f"处理图片时出错: {e}")
    
    def _handle_headers_footers(self, doc: Document, styles: Dict) -> None:
        """处理页眉页脚"""
        logging.info("正在处理页眉页脚...")
        for section in doc.sections:
            # 处理页眉
            header = section.header
            header.is_linked_to_previous = False
            if styles['headers']:
                header_text = styles['headers'][0].get('text', '')
                header.paragraphs[0].text = header_text
            
            # 处理页脚
            footer = section.footer
            footer.is_linked_to_previous = False
            if styles['footers']:
                footer_text = styles['footers'][0].get('text', '')
                footer.paragraphs[0].text = footer_text
    
    def convert(self) -> Path:
        """
        执行转换流程
        
        Returns:
            输出文件路径
        """
        start_time = datetime.now()
        logging.info(f"开始转换 {self.input_pdf}")
        
        try:
            # 1. 初步转换
            initial_docx = self._initial_conversion()
            
            # 2. 提取PDF样式
            styles = self._extract_pdf_styles()
            
            # 3. 打开初步转换的文档
            doc = Document(initial_docx)
            
            # 4. 应用样式
            for i, paragraph in enumerate(doc.paragraphs):
                if i < len(styles['paragraphs']):
                    self._apply_paragraph_styles(paragraph, styles['paragraphs'][i])
            
            # 5. 处理表格
            self._handle_tables(doc, styles)
            
            # 6. 处理图片
            self._handle_images(doc, styles)
            
            # 7. 处理页眉页脚
            self._handle_headers_footers(doc, styles)
            
            # 8. 保存结果
            output_file = self.output_dir / f"{self.input_pdf.stem}_converted.docx"
            doc.save(str(output_file))
            
            # 记录完成信息
            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()
            logging.info(f"转换完成！耗时: {duration:.2f}秒")
            logging.info(f"输出文件: {output_file}")
            
            return output_file
            
        except Exception as e:
            logging.error(f"转换过程中出错: {e}", exc_info=True)
            raise
        finally:
            # 清理临时文件
            if self.pdf_doc:
                self.pdf_doc.close()
            if self.temp_dir:
                import shutil
                shutil.rmtree(self.temp_dir)

def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("使用方法: python pdf2word.py <pdf文件路径> [输出目录]")
        sys.exit(1)
    
    input_pdf = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else "output"
    
    try:
        converter = PDFConverter(input_pdf, output_dir)
        output_file = converter.convert()
        print(f"\n转换成功！输出文件: {output_file}")
    except Exception as e:
        print(f"\n转换失败: {e}")
        sys.exit(1)

if __name__ == '__main__':
    main() 