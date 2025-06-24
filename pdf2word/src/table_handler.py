#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
表格处理模块
用于处理PDF中的表格并转换为Word表格
"""

import logging
from typing import Dict, List, Tuple

import fitz
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.table import _Cell, Table

class TableHandler:
    """表格处理器"""
    
    def __init__(self, pdf_doc: fitz.Document):
        """
        初始化表格处理器
        
        Args:
            pdf_doc: PyMuPDF文档对象
        """
        self.pdf_doc = pdf_doc
    
    def _detect_tables(self, page: fitz.Page) -> List[Dict]:
        """
        检测页面中的表格
        
        Args:
            page: PDF页面对象
        
        Returns:
            表格信息列表
        """
        tables = []
        
        # 使用PyMuPDF的表格检测功能
        tab = page.find_tables()
        if tab.tables:
            for idx, table in enumerate(tab.tables):
                cells = []
                for row in table.cells:
                    row_cells = []
                    for cell in row:
                        if cell is not None:
                            # 提取单元格内容和样式
                            content = page.get_text("text", clip=cell.rect)
                            style = {
                                'rect': cell.rect,
                                'rowspan': cell.rowspan,
                                'colspan': cell.colspan,
                                'text': content.strip()
                            }
                            row_cells.append(style)
                    cells.append(row_cells)
                
                tables.append({
                    'cells': cells,
                    'rect': table.rect,
                    'rows': len(table.cells),
                    'cols': len(table.cells[0]) if table.cells else 0
                })
        
        return tables
    
    def _create_word_table(self, doc: Document, table_info: Dict) -> Table:
        """
        在Word文档中创建表格
        
        Args:
            doc: Word文档对象
            table_info: 表格信息
        
        Returns:
            Word表格对象
        """
        table = doc.add_table(rows=table_info['rows'], 
                            cols=table_info['cols'])
        table.style = 'Table Grid'
        
        # 处理单元格
        for i, row in enumerate(table_info['cells']):
            for j, cell_info in enumerate(row):
                try:
                    cell = table.cell(i, j)
                    
                    # 设置单元格文本
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(cell_info['text'])
                    
                    # 设置字体
                    font = run.font
                    font.size = Pt(10)  # 默认字号
                    
                    # 设置单元格对齐方式
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # 处理合并单元格
                    if cell_info.get('rowspan', 1) > 1 or cell_info.get('colspan', 1) > 1:
                        end_row = i + cell_info.get('rowspan', 1) - 1
                        end_col = j + cell_info.get('colspan', 1) - 1
                        cell.merge(table.cell(end_row, end_col))
                
                except Exception as e:
                    logging.warning(f"处理单元格时出错 ({i},{j}): {e}")
        
        return table
    
    def _set_table_borders(self, table: Table) -> None:
        """设置表格边框"""
        from docx.shared import Pt
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # 设置表格边框
        tbl = table._tbl
        tblPr = tbl.get_or_add_tblPr()
        tblBorders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                             '<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                             '<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                             '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                             '<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                             '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                             '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                             '</w:tblBorders>')
        tblPr.append(tblBorders)
    
    def process_tables(self, doc: Document) -> None:
        """
        处理文档中的所有表格
        
        Args:
            doc: Word文档对象
        """
        logging.info("开始处理表格...")
        
        for page_num in range(len(self.pdf_doc)):
            page = self.pdf_doc[page_num]
            tables = self._detect_tables(page)
            
            for table_info in tables:
                try:
                    # 创建表格
                    table = self._create_word_table(doc, table_info)
                    
                    # 设置边框
                    self._set_table_borders(table)
                    
                    # 添加段落间距
                    doc.add_paragraph()
                    
                except Exception as e:
                    logging.error(f"处理表格时出错: {e}")
        
        logging.info("表格处理完成") 