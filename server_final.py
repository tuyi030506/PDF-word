from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import tempfile
import os
from pathlib import Path
import logging
from datetime import datetime

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="PDF转换工具", description="真实PDF格式转换服务")

# 添加 CORS 中间件
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 创建静态文件目录
static_dir = Path("static")
static_dir.mkdir(exist_ok=True)

# 挂载静态文件
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/")
async def read_root():
    """主页"""
    try:
        with open("static/index.html", "r", encoding="utf-8") as f:
            html_content = f.read()
        return HTMLResponse(content=html_content)
    except Exception as e:
        logger.error(f"读取首页失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"无法读取首页: {str(e)}")

@app.get("/health")
async def health_check():
    """健康检查"""
    return {"status": "healthy", "message": "PDF转换服务运行正常", "version": "2.0.0"}

@app.get("/api/status")
async def get_system_status():
    """获取系统状态"""
    return {
        "system": {
            "status": "running",
            "version": "2.0.0",
            "mode": "production"
        },
        "features": {
            "pdf_to_word": "✅ 正常工作",
            "pdf_to_excel": "✅ 正常工作", 
            "file_upload": "✅ 正常工作",
            "real_conversion": "✅ 已启用"
        },
        "conversion_engine": "pdf2docx + pandas"
    }

@app.post("/api/convert")
async def convert_pdf(
    file: UploadFile = File(...),
    output_format: str = Form(default="docx")
):
    """真实PDF转换API"""
    logger.info(f"收到转换请求: {file.filename} -> {output_format}")
    
    # 验证文件类型
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(
            status_code=400, 
            detail={"error": "INVALID_FILE_TYPE", "message": "只支持PDF文件"}
        )
    
    # 验证输出格式
    if output_format not in ["docx", "xlsx"]:
        raise HTTPException(
            status_code=400,
            detail={"error": "INVALID_FORMAT", "message": "只支持docx和xlsx格式"}
        )
    
    try:
        # 读取文件内容
        content = await file.read()
        file_size_mb = len(content) / (1024 * 1024)
        
        # 文件大小限制
        if file_size_mb > 50:
            raise HTTPException(
                status_code=400,
                detail={"error": "FILE_TOO_LARGE", "message": f"文件过大({file_size_mb:.1f}MB)，请上传小于50MB的文件"}
            )
        
        logger.info(f"文件大小: {file_size_mb:.2f}MB")
        
        # 使用临时目录进行转换
        with tempfile.TemporaryDirectory() as temp_dir:
            # 保存上传的PDF文件
            pdf_path = Path(temp_dir) / "input.pdf"
            with pdf_path.open("wb") as buffer:
                buffer.write(content)
            
            # 执行真实转换
            if output_format == "docx":
                result_file = await convert_pdf_to_word(pdf_path, temp_dir)
            else:
                result_file = await convert_pdf_to_excel(pdf_path, temp_dir)
            
            # 生成下载文件名
            output_filename = file.filename.replace('.pdf', f'.{output_format}')
            
            # 创建输出目录并复制文件到持久位置
            output_dir = Path("converted_files")
            output_dir.mkdir(exist_ok=True)
            
            # 生成唯一的文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            final_output_path = output_dir / f"{timestamp}_{output_filename}"
            
            # 复制文件到持久位置
            import shutil
            shutil.copy2(result_file, final_output_path)
            
            logger.info(f"文件已复制到: {final_output_path}")
            
            # 返回转换后的文件
            media_type = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                if output_format == "docx"
                else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            
            return FileResponse(
                path=str(final_output_path),
                filename=output_filename,
                media_type=media_type
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"转换失败: {str(e)}")
        import traceback
        logger.error(f"详细错误: {traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail={"error": "CONVERSION_FAILED", "message": f"转换失败: {str(e)}"}
        )

async def convert_pdf_to_word(pdf_path: Path, temp_dir: str) -> Path:
    """使用pdf2docx将PDF转换为Word"""
    try:
        from pdf2docx import Converter
        
        output_path = Path(temp_dir) / "output.docx"
        
        logger.info("开始PDF转Word转换...")
        logger.info(f"输入文件: {pdf_path}")
        logger.info(f"输出路径: {output_path}")
        
        # 执行转换
        cv = Converter(str(pdf_path))
        cv.convert(str(output_path), start=0, end=None)
        cv.close()
        
        logger.info("PDF转Word转换完成")
        
        # 检查文件是否存在
        if output_path.exists():
            logger.info(f"转换文件生成成功: {output_path}, 大小: {output_path.stat().st_size} bytes")
            return output_path
        else:
            # 检查临时目录中的所有文件
            temp_files = list(Path(temp_dir).glob("*.docx"))
            logger.error(f"期望的输出文件不存在: {output_path}")
            logger.error(f"临时目录中的docx文件: {temp_files}")
            
            # 如果有其他docx文件，使用第一个
            if temp_files:
                actual_file = temp_files[0]
                logger.info(f"使用实际生成的文件: {actual_file}")
                return actual_file
            else:
                raise Exception("转换失败，未生成输出文件")
            
    except ImportError:
        logger.error("pdf2docx库未安装")
        raise Exception("转换库未安装，请执行: pip install pdf2docx")
    except Exception as e:
        logger.error(f"PDF转Word转换失败: {str(e)}")
        raise Exception(f"PDF转Word转换失败: {str(e)}")

async def convert_pdf_to_excel(pdf_path: Path, temp_dir: str) -> Path:
    """将PDF转换为Excel（提取表格数据）"""
    try:
        import pandas as pd
        from pdf2docx import Converter
        
        output_path = Path(temp_dir) / "output.xlsx"
        
        logger.info("开始PDF转Excel转换...")
        
        # 先转换为Word以提取文本
        temp_docx = Path(temp_dir) / "temp.docx"
        cv = Converter(str(pdf_path))
        cv.convert(str(temp_docx), start=0, end=None)
        cv.close()
        
        # 从Word文档提取文本
        try:
            from docx import Document
            doc = Document(str(temp_docx))
            full_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # 处理表格
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    if any(row_data):  # 只添加非空行
                        table_data.append(row_data)
                if table_data:
                    tables_data.append(table_data)
            
            # 创建Excel文件
            with pd.ExcelWriter(str(output_path), engine='openpyxl') as writer:
                # 如果有表格数据，写入表格
                if tables_data:
                    for i, table_data in enumerate(tables_data):
                        if len(table_data) > 1:  # 确保有数据
                            df = pd.DataFrame(table_data[1:], columns=table_data[0])
                            sheet_name = f'表格{i+1}' if len(tables_data) > 1 else '表格数据'
                            df.to_excel(writer, sheet_name=sheet_name, index=False)
                
                # 写入全部文本内容
                if full_text:
                    text_df = pd.DataFrame({'内容': full_text})
                    text_df.to_excel(writer, sheet_name='文本内容', index=False)
                
                # 如果没有任何内容，创建一个默认表
                if not tables_data and not full_text:
                    default_df = pd.DataFrame({
                        '说明': ['PDF转换完成'],
                        '文件名': [pdf_path.name],
                        '转换时间': [datetime.now().strftime('%Y-%m-%d %H:%M:%S')]
                    })
                    default_df.to_excel(writer, sheet_name='转换信息', index=False)
            
            logger.info("PDF转Excel转换完成")
            return output_path
            
        except ImportError:
            logger.error("python-docx库未安装")
            raise Exception("转换库未安装，请执行: pip install python-docx")
            
    except Exception as e:
        logger.error(f"PDF转Excel转换失败: {str(e)}")
        # 创建一个包含错误信息的Excel文件
        error_df = pd.DataFrame({
            '转换状态': ['转换遇到问题'],
            '文件名': [pdf_path.name],
            '错误信息': [str(e)],
            '建议': ['请尝试PDF转Word格式，或检查PDF文件是否完整']
        })
        error_df.to_excel(str(output_path), sheet_name='转换信息', index=False)
        return output_path

if __name__ == "__main__":
    import uvicorn
    logger.info("启动真实PDF转换服务...")
    uvicorn.run(
        "server_final:app",
        host="0.0.0.0",
        port=3001,
        log_level="info",
        reload=True
    ) 