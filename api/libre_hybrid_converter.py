"""
LibreOffice混合转换器
优先使用LibreOffice，失败时回退到PyPDF2
"""

import logging
from typing import Tuple, Dict, Any
from .libreoffice_converter import LibreOfficeConverter
import PyPDF2
from docx import Document
from docx.shared import Inches
import io
import re

logger = logging.getLogger(__name__)

class LibreHybridConverter:
    """混合转换器：LibreOffice + PyPDF2备用"""
    
    def __init__(self):
        self.libreoffice = LibreOfficeConverter()
        self.stats = {
            "total_conversions": 0,
            "libreoffice_success": 0,
            "pypdf2_fallback": 0,
            "total_failures": 0
        }
    
    def get_status(self) -> Dict[str, Any]:
        """获取转换器状态"""
        libreoffice_status = self.libreoffice.check_installation()
        
        return {
            "libreoffice": libreoffice_status,
            "pypdf2_available": True,  # PyPDF2总是可用的
            "preferred_method": "LibreOffice" if libreoffice_status["installed"] else "PyPDF2",
            "conversion_stats": self.stats.copy()
        }
    
    def convert_pdf_to_docx(self, pdf_content: bytes, filename: str = "document.pdf") -> Tuple[bool, bytes, str]:
        """
        智能PDF到DOCX转换
        1. 优先尝试LibreOffice（高质量）
        2. 失败时回退到PyPDF2（基础质量）
        
        Args:
            pdf_content: PDF文件字节内容
            filename: 原始文件名
            
        Returns:
            (success, docx_content, message)
        """
        self.stats["total_conversions"] += 1
        
        # 尝试LibreOffice转换
        if self.libreoffice.is_available:
            logger.info("Attempting LibreOffice conversion...")
            success, docx_content, message = self.libreoffice.convert_pdf_to_docx(pdf_content, filename)
            
            if success:
                self.stats["libreoffice_success"] += 1
                logger.info("LibreOffice conversion successful")
                return True, docx_content, f"✅ 高质量转换成功 (LibreOffice): {message}"
            else:
                logger.warning(f"LibreOffice conversion failed: {message}")
        else:
            logger.info("LibreOffice not available, using PyPDF2 fallback")
        
        # 回退到PyPDF2转换
        logger.info("Using PyPDF2 fallback conversion...")
        success, docx_content, message = self._convert_with_pypdf2(pdf_content, filename)
        
        if success:
            self.stats["pypdf2_fallback"] += 1
            logger.info("PyPDF2 fallback conversion successful")
            return True, docx_content, f"⚠️  基础转换成功 (PyPDF2): {message}"
        else:
            self.stats["total_failures"] += 1
            logger.error(f"All conversion methods failed: {message}")
            return False, b"", f"❌ 转换失败: {message}"
    
    def _convert_with_pypdf2(self, pdf_content: bytes, filename: str) -> Tuple[bool, bytes, str]:
        """
        使用PyPDF2进行基础PDF转换
        增强版：支持基础格式、段落识别、表格处理
        """
        try:
            # 读取PDF
            pdf_stream = io.BytesIO(pdf_content)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            if len(pdf_reader.pages) == 0:
                return False, b"", "PDF文件没有页面"
            
            # 创建Word文档
            doc = Document()
            
            # 添加标题
            title = doc.add_heading('转换文档', 0)
            title.alignment = 1  # 居中对齐
            
            total_text = ""
            
            # 逐页处理
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    # 提取文本
                    text = page.extract_text()
                    
                    if text.strip():
                        total_text += text + "\n"
                        
                        # 添加页面标题
                        if len(pdf_reader.pages) > 1:
                            page_heading = doc.add_heading(f'第 {page_num} 页', level=2)
                            page_heading.alignment = 0  # 左对齐
                        
                        # 处理文本内容
                        self._process_text_content(doc, text)
                    
                except Exception as e:
                    logger.warning(f"处理第{page_num}页时出错: {e}")
                    # 添加错误页面标记
                    doc.add_paragraph(f"[第{page_num}页处理失败: {str(e)}]")
            
            # 检查是否成功提取到文本
            if not total_text.strip():
                return False, b"", "无法从PDF中提取文本内容"
            
            # 添加转换信息
            doc.add_page_break()
            info_para = doc.add_paragraph()
            info_para.add_run("转换信息:").bold = True
            doc.add_paragraph(f"• 原文件: {filename}")
            doc.add_paragraph(f"• 页数: {len(pdf_reader.pages)}")
            doc.add_paragraph(f"• 转换方式: PyPDF2 基础转换")
            doc.add_paragraph(f"• 字符数: {len(total_text)}")
            doc.add_paragraph("• 注意: 此转换保留了文本内容，但可能丢失原始格式")
            
            # 保存到字节流
            docx_stream = io.BytesIO()
            doc.save(docx_stream)
            docx_content = docx_stream.getvalue()
            
            return True, docx_content, f"成功提取{len(total_text)}个字符，{len(pdf_reader.pages)}页内容"
            
        except Exception as e:
            logger.error(f"PyPDF2转换失败: {e}")
            return False, b"", f"PyPDF2转换错误: {str(e)}"
    
    def _process_text_content(self, doc: Document, text: str):
        """
        处理和格式化文本内容
        包括段落识别、表格检测、列表处理等
        """
        # 分割成行
        lines = text.split('\n')
        current_paragraph = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # 空行 - 结束当前段落
                if current_paragraph:
                    self._add_paragraph(doc, ' '.join(current_paragraph))
                    current_paragraph = []
                continue
            
            # 检测标题（全大写或特殊格式）
            if self._is_likely_heading(line):
                # 先完成当前段落
                if current_paragraph:
                    self._add_paragraph(doc, ' '.join(current_paragraph))
                    current_paragraph = []
                
                # 添加标题
                heading = doc.add_heading(line, level=3)
                heading.alignment = 0
                continue
            
            # 检测列表项
            if self._is_list_item(line):
                # 先完成当前段落
                if current_paragraph:
                    self._add_paragraph(doc, ' '.join(current_paragraph))
                    current_paragraph = []
                
                # 添加列表项
                self._add_list_item(doc, line)
                continue
            
            # 检测表格行（简单检测）
            if self._is_table_row(line):
                # 先完成当前段落
                if current_paragraph:
                    self._add_paragraph(doc, ' '.join(current_paragraph))
                    current_paragraph = []
                
                # 添加表格行（作为段落，因为创建真正的表格比较复杂）
                table_para = doc.add_paragraph(line)
                table_para.style = 'List Bullet'
                continue
            
            # 普通文本行
            current_paragraph.append(line)
            
            # 如果当前段落太长，就结束它
            if len(' '.join(current_paragraph)) > 500:
                self._add_paragraph(doc, ' '.join(current_paragraph))
                current_paragraph = []
        
        # 处理最后的段落
        if current_paragraph:
            self._add_paragraph(doc, ' '.join(current_paragraph))
    
    def _is_likely_heading(self, line: str) -> bool:
        """判断是否可能是标题"""
        # 全大写且不太长
        if line.isupper() and 5 <= len(line) <= 50:
            return True
        
        # 以数字开头的标题
        if re.match(r'^\d+\.?\s+[A-Z]', line):
            return True
        
        # 特殊标题模式
        heading_patterns = [
            r'^第.*章',
            r'^第.*节',
            r'^第.*部分',
            r'^\d+\.\d+',
            r'^[一二三四五六七八九十]+、'
        ]
        
        for pattern in heading_patterns:
            if re.match(pattern, line):
                return True
        
        return False
    
    def _is_list_item(self, line: str) -> bool:
        """判断是否是列表项"""
        list_patterns = [
            r'^\d+\.',  # 1. 2. 3.
            r'^[a-zA-Z]\.',  # a. b. c.
            r'^[•·◦▪▫]',  # 各种项目符号
            r'^[-*+]',  # - * +
            r'^\(\d+\)',  # (1) (2) (3)
        ]
        
        for pattern in list_patterns:
            if re.match(pattern, line):
                return True
        
        return False
    
    def _is_table_row(self, line: str) -> bool:
        """简单判断是否可能是表格行"""
        # 包含多个制表符或连续空格的行
        if '\t' in line or '  ' in line:
            # 且包含多个单词
            words = line.split()
            if len(words) >= 3:
                return True
        
        return False
    
    def _add_paragraph(self, doc: Document, text: str):
        """添加段落"""
        if text.strip():
            para = doc.add_paragraph(text.strip())
            para.alignment = 0  # 左对齐
    
    def _add_list_item(self, doc: Document, text: str):
        """添加列表项"""
        # 移除原有的项目符号
        clean_text = re.sub(r'^[\d+\w\)\.\-\*\+•·◦▪▫\(\)]+\s*', '', text).strip()
        if clean_text:
            para = doc.add_paragraph(clean_text, style='List Bullet')
    
    def get_installation_guide(self) -> Dict[str, Any]:
        """获取安装指南"""
        libreoffice_status = self.libreoffice.check_installation()
        
        if libreoffice_status["installed"]:
            return {
                "status": "ready",
                "message": "LibreOffice已安装，可以进行高质量转换",
                "current_method": "LibreOffice (推荐)",
                "version": libreoffice_status["version"]
            }
        else:
            instructions = self.libreoffice.get_installation_instructions()
            return {
                "status": "needs_installation", 
                "message": "建议安装LibreOffice以获得最佳转换质量",
                "current_method": "PyPDF2 (基础转换)",
                "installation_guide": instructions,
                "fallback": "当前使用PyPDF2进行基础转换"
            }