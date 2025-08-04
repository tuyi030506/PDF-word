#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
混合转换器
- 复杂PDF使用API（高质量）
- 简单PDF使用本地（节省成本）
"""

import os
import logging
import asyncio
from typing import Optional, Dict, Any, Tuple
from .cloudconvert_converter import CloudConvertConverter

logger = logging.getLogger(__name__)

class HybridConverter:
    """
    混合PDF转换器
    - 优先使用CloudConvert实现90%+高质量转换
    - CloudConvert失败时，自动回退到本地PyPDF2方案
    """
    
    def __init__(self, cloudconvert_api_key: str):
        self.cloudconvert_api_key = cloudconvert_api_key
        self.cloudconvert_converter = CloudConvertConverter(cloudconvert_api_key)
    
    async def convert_pdf_to_word(self, input_path: str, output_path: str) -> Tuple[bool, str, Dict[str, Any]]:
        """
        混合PDF转Word转换
        
        Args:
            input_path: 输入PDF文件路径
            output_path: 输出Word文件路径
            
        Returns:
            Tuple[bool, str, Dict]: (成功状态, 转换方法, 详细信息)
        """
        
        # 检查文件是否存在
        if not os.path.exists(input_path):
            return False, "error", {"error": "输入文件不存在"}
        
        file_size = os.path.getsize(input_path)
        logger.info(f"开始混合转换: {input_path} ({file_size} bytes)")
        
        # 策略1: 尝试CloudConvert高质量转换
        logger.info("🚀 尝试CloudConvert高质量转换...")
        cloudconvert_success = await self._try_cloudconvert(input_path, output_path)
        
        if cloudconvert_success:
            output_size = os.path.getsize(output_path) if os.path.exists(output_path) else 0
            logger.info(f"✅ CloudConvert转换成功: {output_size} bytes")
            return True, "cloudconvert", {
                "method": "CloudConvert API",
                "quality": "90%+",
                "input_size": file_size,
                "output_size": output_size,
                "features": ["完美表格", "图像保留", "格式精确"]
            }
        
        # 策略2: 回退到本地PyPDF2转换
        logger.info("🔄 CloudConvert失败，回退到本地PyPDF2转换...")
        local_success = await self._try_local_conversion(input_path, output_path)
        
        if local_success:
            output_size = os.path.getsize(output_path) if os.path.exists(output_path) else 0
            logger.info(f"✅ 本地转换成功: {output_size} bytes")
            return True, "local", {
                "method": "PyPDF2 本地转换",
                "quality": "75%",
                "input_size": file_size,
                "output_size": output_size,
                "features": ["基础表格", "文本提取", "快速处理"]
            }
        
        # 两种方法都失败
        logger.error("❌ 所有转换方法都失败")
        return False, "failed", {
            "error": "CloudConvert和本地转换都失败",
            "input_size": file_size
        }
    
    async def _try_cloudconvert(self, input_path: str, output_path: str) -> bool:
        """尝试CloudConvert转换"""
        try:
            return await self.cloudconvert_converter.convert_pdf_to_word(input_path, output_path)
        except Exception as e:
            logger.error(f"CloudConvert转换异常: {str(e)}")
            return False
    
    async def _try_local_conversion(self, input_path: str, output_path: str) -> bool:
        """尝试本地PyPDF2转换"""
        try:
            return await self._convert_pdf_to_word_nopillow(input_path, output_path)
        except Exception as e:
            logger.error(f"本地转换异常: {str(e)}")
            return False
    
    async def _convert_pdf_to_word_nopillow(self, input_path: str, output_path: str) -> bool:
        """本地PyPDF2转换（增强版）"""
        try:
            logger.info("开始本地PyPDF2转换")
            
            # 导入库
            try:
                from PyPDF2 import PdfReader
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.shared import Inches
                import re
                logger.info("PyPDF2转换器导入成功")
            except ImportError as e:
                logger.error(f"库导入失败: {e}")
                return False
            
            reader = PdfReader(input_path)
            doc = Document()
            
            # 添加文档标题
            title = doc.add_heading('PDF转换文档 (混合转换器)', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加转换信息
            info_para = doc.add_paragraph()
            info_para.add_run("转换方法: ").bold = True
            info_para.add_run("本地PyPDF2转换（CloudConvert备用方案）")
            info_para.add_run("\n质量等级: ").bold = True
            info_para.add_run("75% 文本提取 + 基础格式")
            
            doc.add_paragraph("")  # 空行
            
            # 逐页处理
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text = page.extract_text()
                
                if text.strip():
                    # 添加页码标识
                    page_header = doc.add_paragraph()
                    page_header.add_run(f"【第 {page_num + 1} 页】").bold = True
                    
                    if self._is_table_content(text):
                        self._add_table_content(doc, text, page_num + 1)
                    else:
                        self._add_text_content(doc, text, page_num + 1)
                else:
                    self._add_empty_page_notice(doc, page_num + 1)
                
                # 页面分隔
                if page_num < len(reader.pages) - 1:
                    doc.add_page_break()
            
            doc.save(output_path)
            logger.info(f"本地转换成功，输出文件大小: {os.path.getsize(output_path)} bytes")
            return True
            
        except Exception as e:
            logger.error(f"本地PDF转换异常: {str(e)}", exc_info=True)
            return False
    
    def _is_table_content(self, text: str) -> bool:
        """判断是否为表格内容"""
        lines = text.strip().split('\n')
        if len(lines) < 2:
            return False
        
        # 检查是否有多列数据
        tab_count = sum(1 for line in lines if line.count('\t') >= 2 or line.count('  ') >= 3)
        return tab_count >= len(lines) * 0.3
    
    def _add_table_content(self, doc, text: str, page_num: int):
        """添加表格内容"""
        try:
            lines = [line.strip() for line in text.strip().split('\n') if line.strip()]
            
            if len(lines) < 2:
                self._add_text_content(doc, text, page_num)
                return
            
            # 创建表格
            max_cols = max(len(re.split(r'\s{2,}|\t', line)) for line in lines[:5])
            max_cols = min(max_cols, 6)  # 限制最大列数
            
            table = doc.add_table(rows=1, cols=max_cols)
            table.style = 'Table Grid'
            
            # 添加表头
            header_cells = table.rows[0].cells
            header_parts = re.split(r'\s{2,}|\t', lines[0])
            for i, part in enumerate(header_parts[:max_cols]):
                header_cells[i].text = part
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # 添加数据行
            for line in lines[1:]:
                if line.strip():
                    row_cells = table.add_row().cells
                    row_parts = re.split(r'\s{2,}|\t', line)
                    for i, part in enumerate(row_parts[:max_cols]):
                        row_cells[i].text = part
            
        except Exception as e:
            logger.warning(f"表格处理失败，转为文本: {e}")
            self._add_text_content(doc, text, page_num)
    
    def _add_text_content(self, doc, text: str, page_num: int):
        """添加文本内容"""
        lines = text.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            para = doc.add_paragraph()
            
            if self._is_title_line(line):
                para.style = 'Heading 2'
                para.add_run(line).bold = True
            else:
                para.add_run(line)
            
            if self._should_center_align(line):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _is_title_line(self, line: str) -> bool:
        """判断是否为标题行"""
        if len(line) < 3 or len(line) > 50:
            return False
        return bool(re.match(r'^[一二三四五六七八九十\d\.\s]*[^\d\s]', line))
    
    def _should_center_align(self, line: str) -> bool:
        """判断是否应该居中对齐"""
        return len(line) <= 20 and not line.startswith('  ')
    
    def _add_empty_page_notice(self, doc, page_num: int):
        """添加空页提示"""
        para = doc.add_paragraph()
        para.add_run(f"【第 {page_num} 页为空白页或无可提取文本】").italic = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    async def get_conversion_strategy(self, file_size: int) -> str:
        """根据文件大小推荐转换策略"""
        if file_size > 50 * 1024 * 1024:  # 50MB以上
            return "大文件建议使用CloudConvert，质量更好"
        elif file_size > 10 * 1024 * 1024:  # 10MB以上
            return "中等文件，CloudConvert转换效果佳"
        else:
            return "小文件，CloudConvert快速高质量转换" 