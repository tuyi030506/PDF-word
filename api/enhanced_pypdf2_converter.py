#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
增强版PyPDF2转换器
改善表格识别和格式保持，提升相似度到75-80%
"""

import re
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class EnhancedPyPDF2Converter:
    
    def convert_pdf_to_word(self, input_path: str, output_path: str) -> bool:
        try:
            reader = PdfReader(input_path)
            doc = Document()
            
            # 添加标题
            title = doc.add_heading('PDF转换文档', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text = page.extract_text()
                
                if text.strip():
                    # 分析页面内容类型
                    if self._is_table_content(text):
                        self._add_table_content(doc, text, page_num + 1)
                    else:
                        self._add_text_content(doc, text, page_num + 1)
                else:
                    self._add_empty_page_notice(doc, page_num + 1)
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"转换失败: {e}")
            return False
    
    def _is_table_content(self, text: str) -> bool:
        """判断是否包含表格内容"""
        # 检测表格特征
        lines = text.split('\n')
        
        # 寻找表格行模式
        table_indicators = 0
        for line in lines:
            # 检测多列数据模式
            if len(re.findall(r'\s{3,}', line)) >= 2:  # 多个空格分隔
                table_indicators += 1
            # 检测数字序号
            if re.match(r'^\s*\d+\s+', line):
                table_indicators += 1
        
        return table_indicators >= 3
    
    def _add_table_content(self, doc: Document, text: str, page_num: int):
        """添加表格内容，尝试重建表格结构"""
        doc.add_heading(f'第 {page_num} 页', level=1)
        
        lines = text.split('\n')
        table_lines = []
        
        # 提取表格行
        for line in lines:
            line = line.strip()
            if line and (re.findall(r'\s{3,}', line) or re.match(r'^\d+\s+', line)):
                # 分割列数据
                columns = re.split(r'\s{3,}', line)
                if len(columns) >= 2:
                    table_lines.append(columns)
        
        if table_lines:
            # 确定最大列数
            max_cols = max(len(row) for row in table_lines)
            
            # 创建表格
            table = doc.add_table(rows=len(table_lines), cols=max_cols)
            table.style = 'Table Grid'
            
            # 填充表格数据
            for i, row_data in enumerate(table_lines):
                row = table.rows[i]
                for j, cell_data in enumerate(row_data):
                    if j < len(row.cells):
                        row.cells[j].text = cell_data.strip()
        else:
            # 如果无法构建表格，添加为段落
            self._add_text_content(doc, text, page_num)
    
    def _add_text_content(self, doc: Document, text: str, page_num: int):
        """添加文本内容，保持段落结构"""
        doc.add_heading(f'第 {page_num} 页', level=1)
        
        # 按段落分割
        paragraphs = text.split('\n\n')
        
        for para in paragraphs:
            para = para.strip()
            if para:
                # 检测标题模式
                if self._is_title_line(para):
                    doc.add_heading(para, level=2)
                else:
                    # 处理长段落的换行
                    clean_para = re.sub(r'\n(?!\s*$)', ' ', para)
                    p = doc.add_paragraph(clean_para)
                    
                    # 如果是居中文本，设置居中对齐
                    if self._should_center_align(clean_para):
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _is_title_line(self, text: str) -> bool:
        """判断是否为标题行"""
        text = text.strip()
        # 短文本且包含特定关键词
        if len(text) < 50 and any(keyword in text for keyword in ['说明', '证明', '附件', '第', '章', '节']):
            return True
        return False
    
    def _should_center_align(self, text: str) -> bool:
        """判断是否应该居中对齐"""
        text = text.strip()
        # 短文本、日期、公司名等
        if len(text) < 30 and any(keyword in text for keyword in ['公司', '日期', '说明', '证明']):
            return True
        return False
    
    def _add_empty_page_notice(self, doc: Document, page_num: int):
        """添加空页面提示"""
        doc.add_heading(f'第 {page_num} 页', level=1)
        p = doc.add_paragraph('(此页面无法提取文本，可能包含图像或特殊格式)')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER 