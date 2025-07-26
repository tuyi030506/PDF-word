#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
PDF转换工具 - 无Pillow版本
使用PyPDF2 + python-docx，避免Pillow依赖问题
"""

import os
import tempfile
import logging
import time
from pathlib import Path
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 创建FastAPI应用
app = FastAPI(
    title="PDF转换工具 - 无Pillow版本",
    description="使用PyPDF2的PDF转换服务",
    version="2.6.0"
)

# 添加CORS中间件
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 配置
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

@app.get("/")
async def read_root():
    """主页面 - 无Pillow版本"""
    html_content = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>PDF转换工具</title>
    <meta http-equiv="Content-Security-Policy" content="default-src 'self'; script-src 'self' 'unsafe-inline'; style-src 'self' 'unsafe-inline';">
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; 
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
            min-height: 100vh; 
            display: flex; 
            align-items: center; 
            justify-content: center; 
            padding: 20px;
        }
        .container { 
            background: white; 
            padding: 2rem; 
            border-radius: 20px; 
            box-shadow: 0 20px 40px rgba(0,0,0,0.1); 
            max-width: 500px; 
            width: 100%; 
        }
        h1 { 
            text-align: center; 
            color: #333; 
            margin-bottom: 1rem; 
            font-size: 1.8rem;
        }
        .version-info {
            background: #e8f5e8;
            border: 1px solid #4caf50;
            color: #2e7d32;
            padding: 1rem;
            border-radius: 8px;
            margin-bottom: 1.5rem;
            font-size: 0.9rem;
            text-align: center;
        }
        .warning {
            background: #fff3cd;
            border: 1px solid #ffeaa7;
            color: #856404;
            padding: 1rem;
            border-radius: 8px;
            margin-bottom: 1.5rem;
            font-size: 0.9rem;
        }
        .upload-area { 
            border: 2px dashed #ddd; 
            border-radius: 12px; 
            padding: 2rem; 
            text-align: center; 
            margin-bottom: 1.5rem; 
            cursor: pointer; 
            transition: all 0.3s ease;
        }
        .upload-area:hover { 
            border-color: #667eea; 
            background: #f8f9ff; 
        }
        .file-input { display: none; }
        button { 
            width: 100%; 
            padding: 12px; 
            border: none; 
            border-radius: 8px; 
            font-size: 16px; 
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
            color: white; 
            cursor: pointer; 
            font-weight: 600; 
            transition: transform 0.2s ease;
        }
        button:hover { transform: translateY(-2px); }
        button:disabled { 
            opacity: 0.6; 
            cursor: not-allowed; 
            transform: none;
        }
        .status { 
            margin-top: 1rem; 
            padding: 1rem; 
            border-radius: 8px; 
            text-align: center; 
            display: none; 
        }
        .status.success { background: #d4edda; color: #155724; border: 1px solid #c3e6cb; }
        .status.error { background: #f8d7da; color: #721c24; border: 1px solid #f5c6cb; }
        .status.info { background: #cce7ff; color: #004085; border: 1px solid #b3d7ff; }
    </style>
</head>
<body>
    <div class="container">
        <h1>📄 PDF转换工具</h1>
        
        <div class="version-info">
            <strong>✅ 无Pillow版本 2.6.0</strong><br>
            使用PyPDF2技术，避免依赖问题
        </div>
        
        <div class="warning">
            <strong>⚠️ 功能说明：</strong><br>
            • 文件大小: 最大10MB<br>
            • 处理时间: 最长30秒<br>
            • 支持文本PDF转Word<br>
            • 图像PDF效果有限
        </div>
        
        <form id="uploadForm">
            <div class="upload-area" id="uploadArea">
                <p>📎 点击选择PDF文件</p>
                <input type="file" id="fileInput" class="file-input" accept=".pdf" required>
            </div>
            
            <button type="submit" id="convertBtn">🚀 开始转换</button>
        </form>
        
        <div class="status" id="status"></div>
    </div>

    <script>
        (function() {
            'use strict';
            
            var uploadArea = document.getElementById('uploadArea');
            var fileInput = document.getElementById('fileInput');
            var form = document.getElementById('uploadForm');
            var convertBtn = document.getElementById('convertBtn');
            var status = document.getElementById('status');

            uploadArea.onclick = function() {
                fileInput.click();
            };
            
            fileInput.onchange = function(e) {
                if (e.target.files.length > 0) {
                    uploadArea.innerHTML = '<p>✅ 已选择: ' + e.target.files[0].name + '</p>';
                }
            };

            function showStatus(message, type) {
                status.textContent = message;
                status.className = 'status ' + type;
                status.style.display = 'block';
            }

            form.onsubmit = function(e) {
                e.preventDefault();
                
                var file = fileInput.files[0];
                if (!file) {
                    showStatus('请选择PDF文件', 'error');
                    return;
                }

                if (file.size > 10 * 1024 * 1024) {
                    showStatus('文件大小不能超过10MB', 'error');
                    return;
                }

                convertBtn.disabled = true;
                convertBtn.textContent = '⏳ 转换中...';
                showStatus('正在转换文件，请稍候...', 'info');

                var formData = new FormData();
                formData.append('file', file);

                var xhr = new XMLHttpRequest();
                
                xhr.onreadystatechange = function() {
                    if (xhr.readyState === 4) {
                        convertBtn.disabled = false;
                        convertBtn.textContent = '🚀 开始转换';

                        if (xhr.status === 200) {
                            var blob = new Blob([xhr.response], {
                                type: 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                            });
                            
                            var url = window.URL.createObjectURL(blob);
                            var a = document.createElement('a');
                            a.style.display = 'none';
                            a.href = url;
                            a.download = file.name.replace('.pdf', '.docx');
                            document.body.appendChild(a);
                            a.click();
                            window.URL.revokeObjectURL(url);
                            document.body.removeChild(a);

                            showStatus('✅ 转换完成，文件已下载！', 'success');
                        } else {
                            try {
                                var errorData = JSON.parse(xhr.responseText);
                                showStatus('❌ 转换失败: ' + (errorData.detail || '未知错误'), 'error');
                            } catch (e) {
                                showStatus('❌ 转换失败: 服务器错误', 'error');
                            }
                        }
                    }
                };

                xhr.onerror = function() {
                    convertBtn.disabled = false;
                    convertBtn.textContent = '🚀 开始转换';
                    showStatus('❌ 网络错误，请重试', 'error');
                };

                xhr.open('POST', '/api/convert');
                xhr.responseType = 'arraybuffer';
                xhr.send(formData);
            };
        })();
    </script>
</body>
</html>"""
    return HTMLResponse(content=html_content)

@app.get("/health")
async def health_check():
    """健康检查"""
    return {
        "status": "healthy",
        "message": "PDF转换服务运行正常 (无Pillow版)",
        "version": "2.6.0",
        "environment": "render"
    }

@app.get("/debug")
async def debug_info():
    """系统调试信息"""
    try:
        import sys
        import pkg_resources
        
        # 检查已安装的包
        installed_packages = {}
        for pkg in pkg_resources.working_set:
            installed_packages[pkg.project_name] = pkg.version
        
        # 检查关键依赖
        key_packages = ['fastapi', 'uvicorn', 'PyPDF2', 'python-docx', 'reportlab']
        package_status = {}
        for pkg in key_packages:
            try:
                module_name = pkg.replace('-', '_').replace('PyPDF2', 'PyPDF2')
                __import__(module_name)
                package_status[pkg] = f"✅ {installed_packages.get(pkg, 'unknown')}"
            except ImportError as e:
                package_status[pkg] = f"❌ {str(e)}"
        
        return {
            "python_version": sys.version,
            "working_directory": os.getcwd(),
            "temp_directory": tempfile.gettempdir(),
            "key_packages": package_status,
            "conversion_method": "PyPDF2 + python-docx",
            "pillow_required": False
        }
    except Exception as e:
        logger.error(f"调试信息获取失败: {str(e)}")
        return {"error": str(e)}

@app.post("/api/convert")
async def convert_pdf(file: UploadFile = File(...)):
    """PDF转换API - 无Pillow版本"""
    start_time = time.time()
    
    try:
        logger.info(f"开始处理转换请求: {file.filename}")
        
        # 验证文件
        if not file.filename or not file.filename.lower().endswith('.pdf'):
            logger.error("文件类型验证失败")
            raise HTTPException(status_code=400, detail="只支持PDF文件")
        
        # 读取文件内容
        content = await file.read()
        file_size_mb = len(content) / (1024 * 1024)
        
        logger.info(f"文件信息: {file.filename}, 大小: {file_size_mb:.2f}MB")
        
        # 检查文件大小
        if len(content) > MAX_FILE_SIZE:
            logger.error(f"文件过大: {file_size_mb:.2f}MB")
            raise HTTPException(status_code=400, detail="文件大小超过10MB限制")
        
        if len(content) == 0:
            logger.error("文件为空")
            raise HTTPException(status_code=400, detail="文件为空")
        
        # 创建临时目录
        temp_dir = tempfile.mkdtemp()
        logger.info(f"创建临时目录: {temp_dir}")
        
        try:
            input_path = os.path.join(temp_dir, "input.pdf")
            output_path = os.path.join(temp_dir, "output.docx")
            
            # 写入PDF文件
            with open(input_path, "wb") as f:
                f.write(content)
            
            logger.info(f"PDF文件已写入: {input_path}")
            
            # 执行转换
            success = await convert_pdf_to_word_nopillow(input_path, output_path)
            
            if not success:
                logger.error("转换函数返回失败")
                raise HTTPException(status_code=500, detail="PDF转换失败，请检查文件是否损坏")
            
            # 检查输出文件
            if not os.path.exists(output_path):
                logger.error(f"输出文件不存在: {output_path}")
                raise HTTPException(status_code=500, detail="转换失败，输出文件未生成")
            
            file_size = os.path.getsize(output_path)
            if file_size == 0:
                logger.error("输出文件为空")
                raise HTTPException(status_code=500, detail="转换失败，输出文件为空")
            
            # 生成下载文件名
            base_name = os.path.splitext(file.filename)[0]
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            download_filename = f"{timestamp}_{base_name}.docx"
            
            conversion_time = time.time() - start_time
            logger.info(f"转换完成: {download_filename}, 耗时: {conversion_time:.2f}秒, 大小: {file_size} bytes")
            
            # 返回文件
            return FileResponse(
                path=output_path,
                filename=download_filename,
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
        finally:
            # 延迟清理临时目录
            import threading
            def cleanup():
                import time
                time.sleep(10)
                import shutil
                try:
                    shutil.rmtree(temp_dir)
                    logger.info(f"临时目录已清理: {temp_dir}")
                except Exception as e:
                    logger.warning(f"清理临时目录失败: {e}")
            
            cleanup_thread = threading.Thread(target=cleanup)
            cleanup_thread.daemon = True
            cleanup_thread.start()
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"转换过程中出现异常: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"系统错误: {str(e)}")

async def convert_pdf_to_word_nopillow(input_path: str, output_path: str) -> bool:
    """PDF转Word - 无Pillow版本"""
    try:
        logger.info("开始使用PyPDF2转换")
        
        # 导入库
        try:
            from PyPDF2 import PdfReader
            from docx import Document
            logger.info("PyPDF2和python-docx导入成功")
        except ImportError as e:
            logger.error(f"库导入失败: {e}")
            return False
        
        # 读取PDF
        reader = PdfReader(input_path)
        num_pages = len(reader.pages)
        logger.info(f"PDF页数: {num_pages}")
        
        # 创建Word文档
        doc = Document()
        doc.add_heading('PDF转换文档', 0)
        
        # 提取文本
        for page_num in range(num_pages):
            page = reader.pages[page_num]
            text = page.extract_text()
            
            if text.strip():
                doc.add_heading(f'第 {page_num + 1} 页', level=1)
                doc.add_paragraph(text)
            else:
                doc.add_heading(f'第 {page_num + 1} 页', level=1)
                doc.add_paragraph('(此页面无法提取文本，可能包含图像或特殊格式)')
        
        # 保存文档
        doc.save(output_path)
        
        # 验证输出文件
        if os.path.exists(output_path):
            output_size = os.path.getsize(output_path)
            logger.info(f"转换成功，输出文件大小: {output_size} bytes")
            return output_size > 0
        else:
            logger.error(f"转换后输出文件不存在: {output_path}")
            return False
            
    except Exception as e:
        logger.error(f"PDF转换异常: {str(e)}", exc_info=True)
        return False 